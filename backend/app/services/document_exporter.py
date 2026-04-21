from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.enum.text import WD_UNDERLINE
from typing import List, Dict, Optional, Tuple
from app.core.schemas import Correction
import os


class DocumentExporter:
    """文档导出器 - 生成带修订和批注的文档"""
    
    @staticmethod
    def create_corrected_document(
        original_path: str,
        all_corrections: Dict[str, List[Correction]],
        output_path: str,
        mode: str = 'revised'
    ) -> str:
        """
        创建修正后的文档
        
        Args:
            original_path: 原始文档路径
            all_corrections: 所有修正（按段落 ID）
            output_path: 输出路径
            mode: 导出模式 ('revised' - 修订标记, 'comments' - 批注, 'merged' - 合并)
            
        Returns:
            输出文件路径
        """
        # 加载原始文档
        doc = Document(original_path)
        
        if mode == 'revised':
            doc = DocumentExporter._add_revisions(doc, all_corrections)
        elif mode == 'comments':
            doc = DocumentExporter._add_comments(doc, all_corrections)
        elif mode == 'merged':
            doc = DocumentExporter._merge_corrections(doc, all_corrections)
        
        # 保存文档
        doc.save(output_path)
        
        return output_path
    
    @staticmethod
    def _add_revisions(
        doc: Document,
        all_corrections: Dict[str, List[Correction]]
    ) -> Document:
        """
        添加修订标记
        
        Args:
            doc: Word 文档
            all_corrections: 修正列表
            
        Returns:
            添加了修订标记的文档
        """
        para_index = 0
        
        for seg_id, corrections in all_corrections.items():
            if para_index >= len(doc.paragraphs):
                break
            
            para = doc.paragraphs[para_index]
            
            for correction in corrections:
                if not correction.original or not correction.suggested:
                    continue
                
                # 查找原文本
                if correction.original in para.text:
                    # 标记删除原文
                    DocumentExporter._mark_deletion(para, correction.original)
                    
                    # 标记插入新文
                    DocumentExporter._mark_insertion(para, correction.suggested)
            
            para_index += 1
        
        return doc
    
    @staticmethod
    def _mark_deletion(para, text: str):
        """标记删除文本"""
        for run in para.runs:
            if text in run.text:
                # 清空原文
                run.text = run.text.replace(text, '')
                
                # 添加删除线标记
                run.font.strike = True
                run.font.color.rgb = None  # 使用默认颜色
    
    @staticmethod
    def _mark_insertion(para, text: str):
        """标记插入文本"""
        # 创建新 run 并添加下划线
        new_run = para.add_run(text)
        new_run.font.underline = WD_UNDERLINE.SINGLE
        new_run.font.color.rgb = None  # 使用默认颜色
    
    @staticmethod
    def _add_comments(
        doc: Document,
        all_corrections: Dict[str, List[Correction]]
    ) -> Document:
        """
        添加批注
        
        Args:
            doc: Word 文档
            all_corrections: 修正列表
            
        Returns:
            添加了批注的文档
        """
        para_index = 0
        
        for seg_id, corrections in all_corrections.items():
            if para_index >= len(doc.paragraphs):
                break
            
            para = doc.paragraphs[para_index]
            
            for correction in corrections:
                # 创建批注内容
                comment_text = (
                    f"[{correction.correction_type.upper()}] {correction.reason}\n"
                    f"原文：{correction.original}\n"
                    f"建议：{correction.suggested}"
                )
                
                # 添加批注（简化实现）
                # 完整的批注实现需要更复杂的 XML 操作
                # 这里使用段落末尾添加批注说明的方式
                
                # 在段落末尾添加批注标记
                if para.text.strip():
                    para.add_run(
                        f"\n[批注] {correction.correction_type}: {correction.reason}"
                    ).font.size = docx.shared.Pt(8)
            
            para_index += 1
        
        return doc
    
    @staticmethod
    def _merge_corrections(
        doc: Document,
        all_corrections: Dict[str, List[Correction]]
    ) -> Document:
        """
        合并所有修正，生成最终文档
        
        Args:
            doc: Word 文档
            all_corrections: 修正列表
            
        Returns:
            合并修正后的文档
        """
        para_index = 0
        
        for seg_id, corrections in all_corrections.items():
            if para_index >= len(doc.paragraphs):
                break
            
            para = doc.paragraphs[para_index]
            
            # 获取段落文本
            text = para.text
            
            # 按原文长度降序排序，避免替换时索引偏移
            sorted_corrections = sorted(
                corrections,
                key=lambda x: len(x.original) if x.original else 0,
                reverse=True
            )
            
            # 依次替换
            for correction in sorted_corrections:
                if correction.original and correction.suggested:
                    text = text.replace(
                        correction.original,
                        correction.suggested
                    )
            
            # 更新段落文本
            para.text = text
            
            para_index += 1
        
        return doc
    
    @staticmethod
    def export_revised_file(
        original_path: str,
        output_path: str,
        all_corrections: Dict[str, List[Correction]]
    ) -> str:
        """
        导出修订文档
        
        Args:
            original_path: 原始文档路径
            output_path: 输出路径
            all_corrections: 修正列表
            
        Returns:
            输出文件路径
        """
        return DocumentExporter.create_corrected_document(
            original_path,
            all_corrections,
            output_path,
            mode='revised'
        )
    
    @staticmethod
    def export_comments_file(
        original_path: str,
        output_path: str,
        all_corrections: Dict[str, List[Correction]]
    ) -> str:
        """
        导出批注文档
        
        Args:
            original_path: 原始文档路径
            output_path: 输出路径
            all_corrections: 修正列表
            
        Returns:
            输出文件路径
        """
        return DocumentExporter.create_corrected_document(
            original_path,
            all_corrections,
            output_path,
            mode='comments'
        )
    
    @staticmethod
    def export_merged_file(
        original_path: str,
        output_path: str,
        all_corrections: Dict[str, List[Correction]]
    ) -> str:
        """
        导出合并后的文档（直接应用所有修正）
        
        Args:
            original_path: 原始文档路径
            output_path: 输出路径
            all_corrections: 修正列表
            
        Returns:
            输出文件路径
        """
        return DocumentExporter.create_corrected_document(
            original_path,
            all_corrections,
            output_path,
            mode='merged'
        )
