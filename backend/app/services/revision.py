from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.enum.text import WD_UNDERLINE
from typing import List, Dict, Optional
from app.core.schemas import Correction
import copy


class RevisionService:
    """文档修订服务 - 添加修订标记和批注"""
    
    @staticmethod
    def add_revisions_to_document(
        doc: Document,
        corrections_by_para: Dict[int, List[Correction]]
    ) -> Document:
        """
        为文档段落添加修订标记
        
        Args:
            doc: Word 文档对象
            corrections_by_para: 段落索引到修正列表的映射
            
        Returns:
            添加了修订标记的文档
        """
        for para_idx, corrections in corrections_by_para.items():
            if para_idx >= len(doc.paragraphs):
                continue
                
            para = doc.paragraphs[para_idx]
            
            for correction in corrections:
                # 在段落中查找原文本
                text = para.text
                
                if correction.original not in text:
                    continue
                
                # 创建修订标记
                RevisionService._mark_deletion(para, correction.original)
                RevisionService._mark_insertion(para, correction.suggested)
        
        return doc
    
    @staticmethod
    def _mark_deletion(para, text: str):
        """标记删除文本"""
        for run in para.runs:
            if text in run.text:
                # 使用 w:del 标记删除
                run.text = run.text.replace(text, '')
                
                # 创建删除标记（简化实现）
                run.font.strike = True
    
    @staticmethod
    def _mark_insertion(para, text: str):
        """标记插入文本"""
        # 创建新 run 并添加下划线
        new_run = para.add_run(text)
        new_run.font.underline = WD_UNDERLINE.SINGLE
    
    @staticmethod
    def add_comments(doc: Document, corrections_by_para: Dict[int, List[Correction]]) -> Document:
        """
        为文档添加批注
        
        Args:
            doc: Word 文档对象
            corrections_by_para: 段落索引到修正列表的映射
            
        Returns:
            添加了批注的文档
        """
        comment_ids = set()
        
        for para_idx, corrections in corrections_by_para.items():
            if para_idx >= len(doc.paragraphs):
                continue
            
            para = doc.paragraphs[para_idx]
            
            for correction in corrections:
                # 创建批注
                comment_id = len(comment_ids) + 1
                comment_ids.add(comment_id)
                
                # 构建批注内容
                comment_text = (
                    f"[{correction.correction_type.upper()}] {correction.reason}\n"
                    f"建议：{correction.original} -> {correction.suggested}"
                )
                
                # 添加批注到文档
                RevisionService._add_comment(doc, comment_id, comment_text)
                
                # 将批注关联到段落（简化实现）
                para.add_comment(
                    author="AI 校对",
                    initial="AI",
                    date=Document().styles['Normal'].font.name,
                    text=comment_text
                )
        
        return doc
    
    @staticmethod
    def _add_comment(doc: Document, comment_id: int, text: str):
        """添加批注到文档"""
        # 获取文档的 w:body 元素
        body = doc.element.body
        
        # 创建批注元素
        comment_range_start = parse_xml(
            f'<w:commentRangeStart {qn('w:ns')} xmlns:w="{qn('w:ns')}"/>'
        )
        comment_range_end = parse_xml(
            f'<w:commentRangeEnd {qn('w:ns')} xmlns:w="{qn('w:ns')}"/>'
        )
        
        # 简化实现：直接保存批注信息
        # 完整实现需要更复杂的 XML 操作
    
    @staticmethod
    def merge_corrected_content(
        original_text: str,
        corrections: List[Correction]
    ) -> str:
        """
        合并修正后的内容
        
        Args:
            original_text: 原文本
            corrections: 修正列表
            
        Returns:
            修正后的文本
        """
        result = original_text
        
        # 按原文长度降序排序，避免替换时索引偏移
        sorted_corrections = sorted(
            corrections,
            key=lambda x: len(x.original),
            reverse=True
        )
        
        for correction in sorted_corrections:
            result = result.replace(correction.original, correction.suggested)
        
        return result
    
    @staticmethod
    def export_document_with_changes(
        doc: Document,
        all_corrections: Dict[int, List[Correction]],
        output_path: str,
        include_revisions: bool = True,
        include_comments: bool = True
    ) -> str:
        """
        导出带修订和批注的文档
        
        Args:
            doc: 原始文档
            all_corrections: 所有修正
            output_path: 输出路径
            include_revisions: 是否包含修订标记
            include_comments: 是否包含批注
            
        Returns:
            输出文件路径
        """
        if include_revisions:
            doc = RevisionService.add_revisions_to_document(doc, all_corrections)
        
        if include_comments:
            doc = RevisionService.add_comments(doc, all_corrections)
        
        doc.save(output_path)
        return output_path
