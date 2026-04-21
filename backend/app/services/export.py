from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_UNDERLINE
from typing import List, Dict, Tuple
from app.core.schemas import Correction


class RevisionMarker:
    """Word 修订标记生成器"""
    
    @staticmethod
    def add_revisions(doc: Document, original_text: str, corrected_text: str,
                     corrections: List[Correction]) -> Document:
        """
        在文档中添加修订标记
        
        Args:
            doc: Word 文档对象
            original_text: 原文本
            corrected_text: 修正后的文本
            corrections: 修正列表
            
        Returns:
            添加了修订标记的文档
        """
        # 为每个修正添加修订标记
        for correction in corrections:
            # 找到原文本的位置
            for para in doc.paragraphs:
                if correction.original in para.text:
                    # 删除原文本
                    for run in para.runs:
                        if correction.original in run.text:
                            # 标记删除
                            run.text = run.text.replace(correction.original, "")
                            run.font.strike = True
                            
                            # 添加修正文本（标记为插入）
                            new_run = para.add_run(correction.suggested)
                            new_run.font.underline = WD_UNDERLINE.SINGLE
                            break
        
        return doc
    
    @staticmethod
    def add_comments(doc: Document, corrections: List[Correction]) -> Document:
        """
        在文档中添加批注
        
        Args:
            doc: Word 文档对象
            corrections: 修正列表
            
        Returns:
            添加了批注的文档
        """
        for correction in corrections:
            # 找到原文本位置并添加批注
            for para in doc.paragraphs:
                if correction.original in para.text:
                    # 创建批注
                    comment = doc.add_comment()
                    comment.text = f"[{correction.correction_type.upper()}] {correction.reason}\n建议：{correction.original} -> {correction.suggested}"
                    
                    # 将批注关联到对应的文本范围
                    for run in para.runs:
                        if correction.original in run.text:
                            # 这里简化处理，实际应该精确关联到文本范围
                            break
        
        return doc


class CommentInjector:
    """批注注入器"""
    
    @staticmethod
    def inject_comments(doc: Document, corrections_by_para: Dict[int, List[Correction]]) -> Document:
        """
        为文档段落批量添加批注
        
        Args:
            doc: Word 文档对象
            corrections_by_para: 段落索引到修正列表的映射
            
        Returns:
            添加了批注的文档
        """
        for para_idx, corrections in corrections_by_para.items():
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                
                for correction in corrections:
                    # 创建批注
                    comment = doc.add_comment()
                    comment.text = f"[{correction.correction_type.upper()}] {correction.reason}\n建议：{correction.original} -> {correction.suggested}"
        
        return doc


class DocumentExporter:
    """文档导出器"""
    
    @staticmethod
    def export_with_revisions(doc: Document, output_path: str) -> str:
        """
        导出带修订标记的文档
        
        Args:
            doc: Word 文档对象
            output_path: 输出路径
            
        Returns:
            输出文件路径
        """
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def export_with_comments(doc: Document, output_path: str) -> str:
        """
        导出带批注的文档
        
        Args:
            doc: Word 文档对象
            output_path: 输出路径
            
        Returns:
            输出文件路径
        """
        doc.save(output_path)
        return output_path
