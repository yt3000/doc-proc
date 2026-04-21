from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import List, Tuple
import re


class DocumentParser:
    """文档解析器 - 提取 docx 文档内容"""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        从 docx 文件中提取纯文本内容
        
        Args:
            file_path: docx 文件路径
            
        Returns:
            提取的文本内容
        """
        doc = Document(file_path)
        text_parts = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def extract_paragraphs(file_path: str) -> List[Tuple[str, int, int]]:
        """
        提取文档段落及其位置信息
        
        Args:
            file_path: docx 文件路径
            
        Returns:
            段落列表，每个元素为 (文本，起始索引，结束索引)
        """
        doc = Document(file_path)
        paragraphs = []
        current_index = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                start_index = current_index
                end_index = current_index + len(text)
                paragraphs.append((text, start_index, end_index))
                current_index = end_index + 2  # 两个换行符
            else:
                current_index += len(para.text) + 1
        
        return paragraphs
    
    @staticmethod
    def validate_file(file_path: str) -> bool:
        """
        验证文件是否为有效的 docx 文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否有效
        """
        try:
            Document(file_path)
            return True
        except Exception:
            return False
